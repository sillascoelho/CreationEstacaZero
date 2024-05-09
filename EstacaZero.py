import pandas as pd
import numpy as np
from openpyxl.drawing.image import Image as XLImage
from matplotlib import pyplot as plt 
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from time import sleep
from colorama import init, Fore

""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""

# MÉTODO AOKI VELLOSO

def parametroKeAlfa():

    coeficienteKeAlfa = [
    [1, 1000, 0.014],
    [12, 800, 0.02],
    [13, 700, 0.024],
    [14, 600, 0.03],
    [15, 500, 0.028],
    [2, 400, 0.03],
    [21, 550, 0.022],
    [22, 450, 0.028],
    [23, 230, 0.034],
    [24, 250, 0.03],
    [3, 200, 0.06],
    [31, 350, 0.024],
    [32, 300, 0.028],
    [33, 220, 0.04],
    [34, 330, 0.03]
                        ]
        
    dfParametroAokitable = pd.DataFrame(coeficienteKeAlfa, columns=["Solo","K (kPa)", "Alfa"])

    return dfParametroAokitable


def fatorCorrecaoEstaca(diametro):

    correcaoAoki = [
    ["Franki", 2.5, (2*2.5)],
    ["Metálica", 1.75, (2*1.75)],
    ["Pré-Moldada", (1+diametro/0.8), 2 * (1+diametro/0.8)],
    ["Escavada", 3, (2*3)],
    ["Raiz", 2, (2*2)],
    ["Hélice Contínua", 2, (2*2)],
    ["Ômega", 2, (2*2)]
                    ]

    dfCorrecaoAoki = pd.DataFrame(correcaoAoki, columns=["Estaca", "F1", "F2"])

    return  dfCorrecaoAoki
    

def consultarKeAlfa(solo):

    df = parametroKeAlfa()

    df2 = df.loc[df['Solo'] == solo]

    listacomK = df2['K (kPa)'].tolist()

    listacomAlfa = df2['Alfa'].tolist()

    listaParametroAoki = [listacomK[0], listacomAlfa[0]]

    return listaParametroAoki
    

def consultarF1eF2(estaca, diametro):

    df = fatorCorrecaoEstaca(diametro)

    df2 = df.loc[df['Estaca'] == estaca]

    listacomF1 = df2['F1'].tolist()

    listacomF2 = df2['F2'].tolist()

    listaFatorCorrecao = [listacomF1[0], listacomF2[0]]

    return listaFatorCorrecao


def calculorpAoki(solo, estaca, diametro, nspt):

    valorK = consultarKeAlfa(solo)[0]

    valorF1 = consultarF1eF2(estaca, diametro)[0]

    rp = ( valorK * nspt ) / valorF1

    return rp


def calculorlAoki(solo, estaca, diametro, nspt):

    valorK = consultarKeAlfa(solo)[0]

    valorAlfa = consultarKeAlfa(solo)[1]

    valorF2 = consultarF1eF2(estaca, diametro)[1]

    rl = ( valorK * nspt * valorAlfa ) / valorF2

    return rl


def propGeometricasEstaca(diametro):

    perimetroEstaca = np.pi * diametro

    areaEstaca = ( pow(diametro, 2) * np.pi ) / 4

    resultadosGeometricosEstaca = [diametro, perimetroEstaca, areaEstaca]

    return resultadosGeometricosEstaca


def valoresKAoki(listaSolos):

    listaValoresK = []

    for ts in range(len(listaSolos)):

        listaValoresK.append(consultarKeAlfa(listaSolos[ts])[0])

    return listaValoresK


def valoresAlfaAoki(listaSolos):
    
    listaValoresAlfa = []

    for ts in range(len(listaSolos)):

        listaValoresAlfa.append(round(consultarKeAlfa(listaSolos[ts])[1] * 100, 2))

    return listaValoresAlfa
    

def valoresrpAoki(listaSolos, estaca, diametro, listaNspt):

    listaNspt.append(listaNspt[-1])

    listaValoresrp = []

    for i in range(len(listaSolos)):

        listaValoresrp.append(round(calculorpAoki(listaSolos[i], estaca, diametro, listaNspt[i+1]), 2))

    listaNspt.pop()

    return listaValoresrp


def valoresRpAoki(listaSolos, estaca, diametro, listaNspt):

    areaEstaca = propGeometricasEstaca(diametro)[2]

    listavaloresRp = []

    listavaloresrp = valoresrpAoki(listaSolos, estaca, diametro, listaNspt)

    for i in range(len(listavaloresrp)):

        listavaloresRp.append(round(areaEstaca * listavaloresrp[i], 2))

    return listavaloresRp


def valoresrlAoki(listaSolos, estaca, diametro, listaNspt):
    
    listaValoresrl = []

    for i in range(len(listaSolos)):

        listaValoresrl.append(round(calculorlAoki(listaSolos[i], estaca, diametro, listaNspt[i]), 2))

    return listaValoresrl


def valoresRlAoki(listaSolos, estaca, diametro, listaNspt):

    areaLateralEstaca = propGeometricasEstaca(diametro)[1]

    listavaloresRl = []

    listavaloresrl = valoresrlAoki(listaSolos, estaca, diametro, listaNspt)

    for i in range(len(listavaloresrl)):

        listavaloresRl.append(round((areaLateralEstaca * listavaloresrl[i]), 2))
    
    return listavaloresRl


def valoresAcumuladosRlAoki(listaSolos, estaca, diametro, listaNspt):

    listaValoresRl = valoresRlAoki(listaSolos, estaca, diametro, listaNspt)

    serie = pd.Series(listaValoresRl)

    listaValoresRlAcumulada = round(serie.cumsum(), 2)

    return listaValoresRlAcumulada


def valoresRtAoki(listaSolos, estaca, diametro, listaNspt):

    lateralAcumulado = valoresAcumuladosRlAoki(listaSolos, estaca, diametro, listaNspt)
    
    resistenciaPonta = valoresRpAoki(listaSolos, estaca, diametro, listaNspt)

    resistenciaTotal = []

    for i in range(len(lateralAcumulado)):
        resistenciaTotal.append(round(lateralAcumulado[i] + resistenciaPonta[i], 2))
    
    return resistenciaTotal


def padm6122Aoki(listaSolos, estaca, diametro, listaNspt):

    listaresTotal = valoresRtAoki(listaSolos, estaca, diametro, listaNspt)

    pa6122 = []

    for i in range(len(listaresTotal)):
        pa6122.append(round((listaresTotal[i] / 2), 2))

    return pa6122


def semResPontaEscavadaAoki(listaSolos, estaca, diametro, listaNspt):

    listasemResPontaEscavada = []

    listaAcumuladoRl = valoresAcumuladosRlAoki(listaSolos, estaca, diametro, listaNspt)

    for i in range(len(listaAcumuladoRl)):

        listasemResPontaEscavada.append(round((listaAcumuladoRl[i] / 2), 2))
        
    return listasemResPontaEscavada


def comResPontaEscavadaAoki(listaSolos, estaca, diametro, listaNspt):

    listavaloresRl = valoresAcumuladosRlAoki(listaSolos, estaca, diametro, listaNspt)

    listavaloresRp = valoresRpAoki(listaSolos, estaca, diametro, listaNspt)

    listaMediaRleRl = []

    for i in range(len(listavaloresRl)):

        if listavaloresRp[i] > listavaloresRl[i]:

            listaMediaRleRl.append(round(((0.95 * listavaloresRl[i] + listavaloresRl[i]) / 2), 2))

        else:

            listaMediaRleRl.append(round(((listavaloresRp[i] + listavaloresRl[i]) / 2), 2))

    return listaMediaRleRl


def cotasDeApoio(listaSolos):

    cotasdeApoio = []

    for i in range(len(listaSolos)):

        cotasdeApoio.append(-( i + 1))

    return cotasdeApoio


def resultadosAoki(listaSolos, estaca, diametro, listaNspt):

    resultadosAoki = {
        "Cotas": cotasDeApoio(listaSolos),
        "K (kPa)": valoresKAoki(listaSolos),
        "α (%)": valoresAlfaAoki(listaSolos),
        "rp (kPa)":valoresrpAoki(listaSolos, estaca, diametro, listaNspt),
        "Rp (kN)": valoresRpAoki(listaSolos, estaca, diametro, listaNspt),
        "rl (kPa)": valoresrlAoki(listaSolos, estaca, diametro, listaNspt),
        "Rl (kN)": valoresRlAoki(listaSolos, estaca, diametro, listaNspt),
        "Rl Acum. (kN)": valoresAcumuladosRlAoki(listaSolos, estaca, diametro, listaNspt),
        "Rt (kN)": valoresRtAoki(listaSolos, estaca, diametro, listaNspt),
        "Padm (kN)": padm6122Aoki(listaSolos, estaca, diametro, listaNspt),
        "Padm Escavada S/Res. Ponta (kN)": semResPontaEscavadaAoki(listaSolos, estaca, diametro, listaNspt),
        "Padm Escavada C/Res. Ponta (kN)": comResPontaEscavadaAoki(listaSolos, estaca, diametro, listaNspt)
    }

    dfResultadosAoki = pd.DataFrame(resultadosAoki)
    
    return dfResultadosAoki

   
def plotagemAoki(listaSolos, estaca, diametro, listaNspt, cargaAdmissivel, niveldAgua):

    listaPa = padm6122Aoki(listaSolos, estaca, diametro, listaNspt)

    profundidade = len(listaNspt) + 1

    if profundidade < niveldAgua:

        profundidade = niveldAgua + 1

    listaPaEscavadaSPonta = semResPontaEscavadaAoki(listaSolos, estaca, diametro, listaNspt)

    listaPaEscavaCPonta = comResPontaEscavadaAoki(listaSolos, estaca, diametro, listaNspt)

    fig, axs = plt.subplots(2, 2, figsize=(12, 9))

    axs[0, 0].plot(listaNspt, cotasDeApoio(listaNspt), label="Nº de Golpes / Camada", color="red")
    axs[0, 0].axhline(y = - niveldAgua, label="Nível d'água", ls="--", color="blue")
    axs[0, 0].set_title("Gráfico da Penetração", fontsize="12", fontweight="bold")
    axs[0, 0].set_xlabel("Número de Golpes", fontsize="10")
    axs[0, 0].set_ylabel("Cotas de Apoio (m)", fontsize="10")
    axs[0, 0].set_yticks(np.arange(0, - profundidade, -1))
    axs[0, 0].grid(True)
    axs[0, 0].legend()

    axs[0, 1].plot(listaPa, cotasDeApoio(listaNspt), label="Carga Admissível / Camada", color="purple")
    axs[0, 1].set_title("Carga Admissível - Aoki Velloso", fontsize="12", fontweight="bold")
    axs[0, 1].axvline(x = cargaAdmissivel, label="Carga Admissível Esperada", ls="--", color="blue")
    axs[0, 1].set_xlabel("Carga Admissível (kN)", fontsize="10")
    axs[0, 1].set_yticks(np.arange(0, - profundidade, -1))
    axs[0, 1].grid(True)
    axs[0, 1].legend()

    axs[1, 0].plot(listaPaEscavadaSPonta, cotasDeApoio(listaNspt), label="Carga Admissível / Camada", color="brown")
    axs[1, 0].set_title("Estaca Escavada S/ Res. Ponta - Aoki Velloso", fontsize="12", fontweight="bold")
    axs[1, 0].axvline(x = cargaAdmissivel, label="Carga Admissível Esperada", ls="--", color="blue")
    axs[1, 0].set_xlabel("Carga Admissível (kN)", fontsize="10")
    axs[1, 0].set_ylabel("Cotas de Apoio (m)", fontsize="10")
    axs[1, 0].set_yticks(np.arange(0, - profundidade, -1))
    axs[1, 0].grid(True)
    axs[1, 0].legend()

    axs[1, 1].plot(listaPaEscavaCPonta, cotasDeApoio(listaNspt), label="Carga Admissível / Camada", color="brown")
    axs[1, 1].set_title("Estaca Escavada C/ Res. Ponta - Aoki Velloso", fontsize="12", fontweight="bold")
    axs[1, 1].axvline(x = cargaAdmissivel, label="Carga Admissível Esperada", ls="--", color="blue")
    axs[1, 1].set_xlabel("Carga Admissível (kN)", fontsize="10")
    axs[1, 1].set_yticks(np.arange(0, - profundidade, -1))
    axs[1, 1].grid(True)
    axs[1, 1].legend()
    
    plt.tight_layout()

    imagemAoki = BytesIO() 

    plt.savefig(imagemAoki, format='png')
    
    imagemAoki.seek(0)
    
    return imagemAoki

""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""

#MÉTODO DECOURT QUARESMA

def parametroC():

    coeficienteC = [
    [1, 400],
    [12, 400],
    [13, 400],
    [14, 400],
    [15, 400],
    [2, 200],
    [21, 250],
    [22, 200],
    [23, 200],
    [24, 200],
    [3, 120],
    [31, 120],
    [32, 120],
    [33, 120],
    [34, 120]
                    ]
        
    dfParametroC = pd.DataFrame(coeficienteC, columns=["Solo","C (kPa)"])

    return dfParametroC


def parametroAlfaeBeta(estaca):

    if estaca == "Escavada":
        
        coeficienteAlfaeBeta = [
        [1, 0.5, 0.5],
        [12, 0.5, 0.5],
        [13, 0.5, 0.5],
        [14, 0.5, 0.5],
        [15, 0.5, 0.5],
        [2, 0.6, 0.65],
        [21, 0.6, 0.65],
        [22, 0.6, 0.65],
        [23, 0.6, 0.65],
        [24, 0.6, 0.65],
        [3, 0.85, 0.8],
        [31, 0.85, 0.8],
        [32, 0.85, 0.8],
        [33, 0.85, 0.8],
        [34, 0.85, 0.8]
                            ]
    
    elif estaca == "Hélice Contínua":

        coeficienteAlfaeBeta = [
        [1, 0.3, 1],
        [12, 0.3, 1],
        [13, 0.3, 1],
        [14, 0.3, 1],
        [15, 0.3, 1],
        [2, 0.3, 1],
        [21, 0.3, 1],
        [22, 0.3, 1],
        [23, 0.3, 1],
        [24, 0.3, 1],
        [3, 0.3, 1],
        [31, 0.3, 1],
        [32, 0.3, 1],
        [33, 0.3, 1],
        [34, 0.3, 1]
                            ]
    
    elif estaca == "Raiz":
        coeficienteAlfaeBeta = [
        [1, 0.5, 1.5],
        [12, 0.5, 1.5],
        [13, 0.5, 1.5],
        [14, 0.5, 1.5],
        [15, 0.5, 1.5],
        [2, 0.6, 1.5],
        [21, 0.6, 1.5],
        [22, 0.6, 1.5],
        [23, 0.6, 1.5],
        [24, 0.6, 1.5],
        [3, 0.85, 1.5],
        [31, 0.85, 1.5],
        [32, 0.85, 1.5],
        [33, 0.85, 1.5],
        [34, 0.85, 1.5]
                            ]
        
    else:
        coeficienteAlfaeBeta = [
        [1, 1, 1],
        [12, 1, 1],
        [13, 1, 1],
        [14, 1, 1],
        [15, 1, 1],
        [2, 1, 1],
        [21, 1, 1],
        [22, 1, 1],
        [23, 1, 1],
        [24, 1, 1],
        [3, 1, 1],
        [31, 1, 1],
        [32, 1, 1],
        [33, 1, 1],
        [34, 1, 1]
                            ]
       
    dfParametroDQtable = pd.DataFrame(coeficienteAlfaeBeta, columns=["Solo","α", "β"])

    return dfParametroDQtable

def consultaParametroC(solo):

    df = parametroC()

    df2 = df.loc[df['Solo'] == solo]

    listacomC = df2['C (kPa)'].tolist()

    return listacomC

def consultaAlfaeBeta(solo, estaca):

    df = parametroAlfaeBeta(estaca)
    
    df2 = df.loc[df['Solo'] == solo]

    listacomAlfa = df2['α'].tolist()

    listacomBeta = df2['β'].tolist()

    listaParametroDQ = [listacomAlfa[0], listacomBeta[0]]

    return listaParametroDQ


def valoresC(listaSolos):

    listaValoresC = []

    for ts in range(len(listaSolos)):

        listaValoresC.append(consultaParametroC(listaSolos[ts])[0])

    return listaValoresC


def valoresBeta(listaSolos, estaca):

    listaValoresBeta = []

    for ts in range(len(listaSolos)):

        listaValoresBeta.append(consultaAlfaeBeta(listaSolos[ts], estaca)[1])

    return listaValoresBeta


def valoresAlfaDQ(listaSolos, estaca):

    listaValoresAlfa = []

    for ts in range(len(listaSolos)):

        listaValoresAlfa.append(consultaAlfaeBeta(listaSolos[ts], estaca)[0])

    return listaValoresAlfa


def nsptp(listaSolos, listaNspt): 

    listaValoresMedioNspt = []

    listaNspt.append(listaNspt[-1])

    for i in range(len(listaSolos)):

        if i == 0:
            nsptMedio =  round(( listaNspt[i] + listaNspt[i] + listaNspt[i + 1] ) /  3, 2)
            listaValoresMedioNspt.append(nsptMedio)

        else:
            nsptMedio = round(( listaNspt[i - 1] + listaNspt[i] + listaNspt[i + 1] ) /  3, 2)
            listaValoresMedioNspt.append(nsptMedio)

    listaNspt.pop()

    listaValoresMedioNspt.pop(0)
    listaValoresMedioNspt.append(listaValoresMedioNspt[-1])

    return listaValoresMedioNspt


def valoresrpDQ(listaSolos, listaNspt):

    listaValoresMediosNspt = nsptp(listaSolos, listaNspt)

    listaValoresC = valoresC(listaSolos)

    listaValoresrp = []

    for i in range(len(listaSolos)):

        listaValoresrp.append(round(listaValoresMediosNspt[i] * listaValoresC[i], 2))

    return listaValoresrp


def valoresRpDQ(listaSolos, estaca, diametro, listaNspt):

    listaValoresRp = []

    rp = valoresrpDQ(listaSolos, listaNspt)

    alfa = valoresAlfaDQ(listaSolos, estaca)

    area = propGeometricasEstaca(diametro)[2]
    
    for i in range(len(listaSolos)):

        listaValoresRp.append(round(rp[i] * alfa[i] * area, 2))

    return listaValoresRp


def valoresrlDQ(listaSolos, listaNspt):

    listarl = []

    for i in range(len(listaSolos)):

        rl =  round(10 * ((listaNspt[i] / 3) + 1), 2)

        listarl.append(rl)

    return listarl


def nsptl(listaSolos, estaca, listaNspt):

    if estaca == "Pré-Moldada":

        listaNsptL = []

        cont = 0

        for i in range(len(listaSolos)):  

            if i == 0:
            
                listaNsptL.append(3)

            cont = cont + 1

            media = round(sum(listaNspt[0:i+1]) / cont, 2)
            
            if media < 3:

                listaNsptL.append(3)

            elif media > 50:

                listaNsptL.append(50)

            else:

                listaNsptL.append(media)

    elif estaca == "Metálica":

        listaNsptL = []

        cont = 0

        for i in range(len(listaSolos)):  

            if i == 0:
            
                listaNsptL.append(3)

            cont = cont + 1

            media = round(sum(listaNspt[0:i+1]) / cont, 2)
            
            if media < 3:

                listaNsptL.append(3)

            elif media > 50:

                listaNsptL.append(50)

            else:

                listaNsptL.append(media)

    elif estaca == "Franki":

        listaNsptL = []

        cont = 0

        for i in range(len(listaSolos)):  

            if i == 0:
            
                listaNsptL.append(3)

            cont = cont + 1

            media = round(sum(listaNspt[0:i+1]) / cont, 2)
            
            if media < 3:

                listaNsptL.append(3)

            elif media > 50:

                listaNsptL.append(50)

            else:

                listaNsptL.append(media)
    else:

        listaNsptL = []

        cont = 0

        for i in range(len(listaSolos)):  

            if i == 0:
            
                listaNsptL.append(3)

            cont = cont + 1

            media = round(sum(listaNspt[0:i+1]) / cont, 2)
            
            if media < 3:

                listaNsptL.append(3)

            elif media > 15:

                listaNsptL.append(15)

            else:

                listaNsptL.append(media)

    listaNsptL.pop()

    return listaNsptL


def valoresNsptl(listaSolos, estaca, listaNspt):

    listaNsptL = []

    nsptL = nsptl(listaSolos, estaca, listaNspt)

    for i in range(len(listaSolos)):

        listaNsptL.append(nsptL[i])

    return listaNsptL


def valoresrlcomNsptL(listaSolos, estaca, listaNspt):
    
    listarlcomNsptL = []

    listaNsptL = nsptl(listaSolos, estaca, listaNspt)

    rl = valoresrlDQ(listaSolos, listaNsptL)

    for i in range(len(listaSolos)):

        if i == 0:

            listarlcomNsptL.append(0)

        else:

            listarlcomNsptL.append(rl[i])

    return listarlcomNsptL


def valoresRlDQ(listaSolos, estaca, diametro, listaNspt):

    listaValoresRl = []

    rl = valoresrlcomNsptL(listaSolos, estaca, listaNspt)

    beta = valoresBeta(listaSolos, estaca)

    perimetro = propGeometricasEstaca(diametro)[1]
    
    for i in range(len(listaSolos)):

        listaValoresRl.append(round(rl[i] * beta[i] * perimetro, 2))

    return listaValoresRl


def valoresAcumuladosRlDQ(listaSolos, estaca, diametro, listaNspt):

    listaValoresRl = valoresRlDQ(listaSolos, estaca, diametro, listaNspt)

    serie = pd.Series(listaValoresRl)

    listaValoresRlAcumulada = round(serie.cumsum(), 2)

    return listaValoresRlAcumulada


def valoresRtDQ(listaSolos, estaca, diametro, listaNspt):

    lateralAcumulado = valoresAcumuladosRlDQ(listaSolos, estaca, diametro, listaNspt)
    
    resistenciaPonta = valoresRpDQ(listaSolos, estaca, diametro, listaNspt)

    resistenciaTotal = []

    for i in range(len(lateralAcumulado)):
        resistenciaTotal.append(round(lateralAcumulado[i] + resistenciaPonta[i], 2))
    
    return resistenciaTotal


def padm6122DQ(listaSolos, estaca, diametro, listaNspt):

    listaresTotal = valoresRtDQ(listaSolos, estaca, diametro, listaNspt)

    pa6122 = []

    for i in range(len(listaresTotal)):

        pa6122.append(round((listaresTotal[i] / 2), 2))

    return pa6122


def semResPontaEscavadaDQ(listaSolos, estaca, diametro, listaNspt):

    listasemResPontaEscavada = []

    listaAcumuladoRl = valoresAcumuladosRlDQ(listaSolos, estaca, diametro, listaNspt)

    for i in range(len(listaAcumuladoRl)):

        listasemResPontaEscavada.append(round((listaAcumuladoRl[i] / 2), 2))
        
    return listasemResPontaEscavada


def comResPontaEscavadaDQ(listaSolos, estaca, diametro, listaNspt):

    listavaloresRl = valoresAcumuladosRlDQ(listaSolos, estaca, diametro, listaNspt)
    
    listavaloresRp = valoresRpDQ(listaSolos, estaca, diametro, listaNspt)

    listaMediaRleRl = []

    for i in range(len(listavaloresRl)):

        if listavaloresRp[i] > listavaloresRl[i]:

            listaMediaRleRl.append(round(((0.95 * listavaloresRl[i] + listavaloresRl[i]) / 2), 2))

        else:

            listaMediaRleRl.append(round((( listavaloresRp[i] + listavaloresRl[i]) / 2), 2))

    return listaMediaRleRl


def valoresNspt(listaSolos, listaNspt):

    listaNspts = []

    for i in range(len(listaSolos)):
        
        listaNspts.append(listaNspt[i])

    return listaNspts


def listanstp(listaSolos, listaNspt):

    listaNsptP = []

    listanstp = nsptp(listaSolos, listaNspt)

    for i in range(len(listaSolos)):

        listaNsptP.append(listanstp[i])

    return listaNsptP


def resultadosDQ(listaSolos, estaca, diametro, listaNspt):

    resultadosDQ = {
    "Cotas": cotasDeApoio(listaSolos),
    "Nsptl":valoresNsptl(listaSolos, estaca, listaNspt),
    "Nsptp":listanstp(listaSolos, listaNspt),
    "C (kPa)": valoresC(listaSolos),
    "α ": valoresAlfaDQ(listaSolos, estaca),
    "β ": valoresBeta(listaSolos, estaca),
    "rp (kPa)": valoresrpDQ(listaSolos, listaNspt),
    "Rp (kN)": valoresRpDQ(listaSolos, estaca, diametro, listaNspt),
    "rl (kPa)": valoresrlcomNsptL(listaSolos, estaca, listaNspt),
    "Rl (kN)": valoresRlDQ(listaSolos, estaca, diametro, listaNspt),
    "Rl Acum. (kN)": valoresAcumuladosRlDQ(listaSolos, estaca, diametro, listaNspt),
    "Rt (kN)": valoresRtDQ(listaSolos, estaca, diametro, listaNspt),
    "Padm (kN)": padm6122DQ(listaSolos, estaca, diametro, listaNspt),
    "Padm Escavada S/Res. Ponta (kN)": semResPontaEscavadaDQ(listaSolos, estaca, diametro, listaNspt),
    "Padm Escavada C/Res. Ponta (kN)": comResPontaEscavadaDQ(listaSolos, estaca, diametro, listaNspt)
    }
    
    dfResultadosDQ = pd.DataFrame(resultadosDQ)

    return dfResultadosDQ


def plotagemDQ(listaSolos, estaca, diametro, listaNspt, cargaAdmissivel, niveldAgua):
   
    listaPa = padm6122DQ(listaSolos, estaca, diametro, listaNspt)

    profundidade = len(listaNspt) + 1

    if profundidade < niveldAgua:

        profundidade = niveldAgua + 1

    listaPaEscavadaSPonta = semResPontaEscavadaDQ(listaSolos, estaca, diametro, listaNspt)

    listaPaEscavaCPonta = comResPontaEscavadaDQ(listaSolos, estaca, diametro, listaNspt)

    fig, axs = plt.subplots(2, 2, figsize=(12, 9))

    axs[0, 0].plot(listaNspt, cotasDeApoio(listaNspt), label="Nº de Golpes / Camada", color="red")
    axs[0, 0].axhline(y = - niveldAgua, label="Nível d'água", ls="--", color="blue")
    axs[0, 0].set_title("Gráfico da Penetração", fontsize="12", fontweight="bold")
    axs[0, 0].set_xlabel("Número de Golpes", fontsize="10")
    axs[0, 0].set_ylabel("Cotas de Apoio (m)", fontsize="10")
    axs[0, 0].set_yticks(np.arange(0, - profundidade, -1))
    axs[0, 0].grid(True)
    axs[0, 0].legend()

    axs[0, 1].plot(listaPa, cotasDeApoio(listaNspt), label="Carga Admissível / Camada", color="purple")
    axs[0, 1].set_title("Carga Admissível - Decourt Quaresma", fontsize="12", fontweight="bold")
    axs[0, 1].axvline(x = cargaAdmissivel, label="Carga Admissível Esperada", ls="--", color="blue")
    axs[0, 1].set_xlabel("Carga Admissível (kN)", fontsize="10")
    axs[0, 1].set_yticks(np.arange(0, - profundidade, -1))
    axs[0, 1].grid(True)
    axs[0, 1].legend()

    axs[1, 0].plot(listaPaEscavadaSPonta, cotasDeApoio(listaNspt), label="Carga Admissível / Camada", color="brown")
    axs[1, 0].set_title("Estaca Escavada S/ Res. Ponta - Decourt Quaresma", fontsize="12", fontweight="bold")
    axs[1, 0].axvline(x = cargaAdmissivel, label="Carga Admissível Esperada", ls="--", color="blue")
    axs[1, 0].set_xlabel("Carga Admissível (kN)", fontsize="10")
    axs[1, 0].set_ylabel("Cotas de Apoio (m)", fontsize="10")
    axs[1, 0].set_yticks(np.arange(0, - profundidade, -1))
    axs[1, 0].grid(True)
    axs[1, 0].legend()

    axs[1, 1].plot(listaPaEscavaCPonta, cotasDeApoio(listaNspt), label="Carga Admissível / Camada", color="brown")
    axs[1, 1].set_title("Estaca Escavada C/ Res. Ponta - Decourt Quaresma", fontsize="12", fontweight="bold")
    axs[1, 1].axvline(x = cargaAdmissivel, label="Carga Admissível Esperada", ls="--", color="blue")
    axs[1, 1].set_xlabel("Carga Admissível (kN)", fontsize="10")
    axs[1, 1].set_yticks(np.arange(0, - profundidade, -1))
    axs[1, 1].grid(True)
    axs[1, 1].legend()
    
    plt.tight_layout()

    imagemDQ = BytesIO() 

    plt.savefig(imagemDQ, format='png')
    
    imagemDQ.seek(0)

    return imagemDQ

""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""

#COMPARATIVO ENTRE OS MÉTODOS

def comparar(listaSolos, estaca, diametro, listaNspt):

    padm_AokiVelloso = padm6122Aoki(listaSolos, estaca, diametro, listaNspt)

    padm_DecourtQuaresma = padm6122DQ(listaSolos, estaca, diametro, listaNspt)

    menorAokiouDecourt = []

    for i in range(len(listaSolos)):

        if padm_AokiVelloso[i] < padm_DecourtQuaresma[i]:

            menorAokiouDecourt.append(padm_AokiVelloso[i])

        else:

            menorAokiouDecourt.append(padm_DecourtQuaresma[i])

    return menorAokiouDecourt


def comparativoAokieDecourt(listaSolos, estaca, diametro, listaNspt):

    padm_AokiVelloso = padm6122Aoki(listaSolos, estaca, diametro, listaNspt)

    padm_DecourtQuaresma = padm6122DQ(listaSolos, estaca, diametro, listaNspt)

    comparativo = {
        "Cotas": cotasDeApoio(listaSolos),
        "Padm - Aoki Velloso (kN)": padm_AokiVelloso,
        "Padm - Decourt Quaresma (kN)": padm_DecourtQuaresma,
        "Menor valor entre os métodos (kN)": comparar(listaSolos, estaca, diametro, listaNspt)
    }

    dfCompararAokieDecourt = pd.DataFrame(comparativo)

    return dfCompararAokieDecourt


def plotCompararAokieDecourt(listaSolos, estaca, diametro, listaNspt, cargaAdmissivel, niveldAgua):
    
    padm_AokiVelloso = padm6122Aoki(listaSolos, estaca, diametro, listaNspt)

    padm_DecourtQuaresma = padm6122DQ(listaSolos, estaca, diametro, listaNspt)

    profundidade = len(listaNspt) + 1

    if profundidade < niveldAgua:

        profundidade = niveldAgua + 1

    fig, axs = plt.subplots(1, 2, figsize=(11, 9))

    axs[0].plot(padm_AokiVelloso, cotasDeApoio(listaSolos), label="Aoki Velloso", color="red", marker="o")
    axs[0].plot(padm_DecourtQuaresma, cotasDeApoio(listaSolos), label="Decourt Quaresma", color="orange", marker="s")
    axs[0].axvline(x = cargaAdmissivel, label="Carga Admissível Esperada", ls="--", color="purple")
    axs[0].set_title("Carga Admissível - Aoki Velloso e Decourt Quaresma", fontsize="12", fontweight="bold")
    axs[0].set_xlabel("Carga Admissível (kN)", fontsize="10")
    axs[0].set_ylabel("Cotas de Apoio (m)", fontsize="10")
    axs[0].set_yticks(np.arange(0, - profundidade, -1))
    axs[0].grid(True)
    axs[0].legend()

    axs[1].plot(comparar(listaSolos, estaca, diametro, listaNspt), cotasDeApoio(listaSolos), label="Menor entre Aoki V. e Decourt Q.", color="blue", marker="X")
    axs[1].axvline(x = cargaAdmissivel, label="Carga Admissível Esperada", ls="--", color="purple")
    axs[1].set_title("Carga Admissível - Considerando o menor valor", fontsize="12", fontweight="bold")
    axs[1].set_xlabel("Carga Admissível (kN)", fontsize="10")
    axs[1].set_yticks(np.arange(0, - profundidade, -1))
    axs[1].grid(True)
    axs[1].legend()

    plt.tight_layout()

    imagemAokieDecourt = BytesIO() 

    plt.savefig(imagemAokieDecourt, format='png')
    
    imagemAokieDecourt.seek(0)
    
    return imagemAokieDecourt

""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""

#EXPORTAR OS DADOS PARA O EXCEL

def excelExport(listaSolos, estaca, diametro, listaNspt, cargaAdmissivel, niveldAgua, fileName):

    with pd.ExcelWriter(fileName + '.xlsx', engine='openpyxl') as writer:

        resultadosAoki(listaSolos, estaca, diametro, listaNspt).to_excel(writer, sheet_name='Resultados Aoki Velloso', index=False)

        imgAoki = plotagemAoki(listaSolos, estaca, diametro, listaNspt, cargaAdmissivel, niveldAgua)
        wb = writer.book
        ws = wb['Resultados Aoki Velloso']
        ws.add_image(XLImage(imgAoki), 'M1')


        resultadosDQ(listaSolos, estaca, diametro, listaNspt).to_excel(writer, sheet_name='Resultados Decourt Quaresma', index=False)
        
        imgDQ = plotagemDQ(listaSolos, estaca, diametro, listaNspt, cargaAdmissivel, niveldAgua)
        wb = writer.book
        ws = wb['Resultados Decourt Quaresma']
        ws.add_image(XLImage(imgDQ), 'Q1')


        comparativoAokieDecourt(listaSolos, estaca, diametro, listaNspt).to_excel(writer, sheet_name='Comparativo entre os Métodos', index= False)

        imgAokieDecourt = plotCompararAokieDecourt(listaSolos, estaca, diametro, listaNspt, cargaAdmissivel, niveldAgua)
        wb = writer.book
        ws = wb['Comparativo entre os Métodos']
        ws.add_image(XLImage(imgAokieDecourt), 'E1')

    return

def wordExport(listaSolos, estaca, diametro, listaNspt, cargaAdmissivel, niveldAgua, fileName):

    documento = Document()

    #TÍTULO
    titulo = documento.add_heading("Resultados Aoki Velloso e Decourt Quaresma", 0)

    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    titulo.style.font.size = Pt(20)

    titulo.style.font.bold = True

    #DADOS

    dados = documento.add_heading("Dados de Entrada", level=1)

    dados.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    paragrafo = documento.add_paragraph(f"""
    Lista Nspt: {listaNspt}
    
    Tipo de Estaca: {estaca}

    Diâmetro da Estaca: {diametro*100} cm

    Carga Admissível Esperada: {cargaAdmissivel} kN

    Nível d'água: {niveldAgua} m""")

    paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    paragrafo.style.font.size = Pt(12)

    documento.add_page_break()

    """"""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""
    
    #GRÁFICO AOKI VELLOSO

    tituloAokiVelloso = documento.add_heading('Resultados Aoki Velloso', level=1)

    space = documento.add_paragraph('')

    imagemAoki = plotagemAoki(listaSolos, estaca, diametro, listaNspt, cargaAdmissivel, niveldAgua)

    figura1nome = documento.add_paragraph("Figura 1: Gráfico dos resultados pelo método de Aoki Velloso.")

    figura1nome.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    figura1 = documento.add_picture(imagemAoki, width=Cm(15), height=Cm(15))

    figura1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    documento.add_page_break()

    #TABELA AOKI VELLOSO

    dfAoki = resultadosAoki(listaSolos, estaca, diametro, listaNspt)

    tabela1nome = documento.add_paragraph("Tabela 1: Resultados pelo método de Aoki Velloso.")

    tabela1nome.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    linhas = dfAoki.shape[0] + 1

    colunas = dfAoki.shape[1]

    tabelaAoki = documento.add_table(rows=linhas, cols=colunas, style="Light Grid Accent 1")

    for numerocoluna, indicecoluna in enumerate(dfAoki):

        for numerolinha in range(linhas):

            celula = tabelaAoki.cell(numerolinha, numerocoluna)

            if numerolinha == 0:

                celula.text = dfAoki.columns[numerocoluna]

                for paragraph in celula.paragraphs:

                    for run in paragraph.runs:

                        run.font.size = Pt(8)
            else:

                celula.text = str(dfAoki[indicecoluna][numerolinha-1])
                
                for paragraph in celula.paragraphs:

                    for run in paragraph.runs:

                        run.font.size = Pt(8)

    tabelaAoki.autofit = True

    tabelaAoki.alignment = WD_TABLE_ALIGNMENT.CENTER

    documento.add_page_break()

    """"""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""

    #GRÁFICO DECOURT QUARESMA

    tituloDecourtQuaresma = documento.add_heading('Resultados Decourt Quaresma', level=1)

    space = documento.add_paragraph('')

    imagemDQ = plotagemDQ(listaSolos, estaca, diametro, listaNspt, cargaAdmissivel, niveldAgua)

    figura2nome = documento.add_paragraph("Figura 2: Gráfico dos resultados pelo método de Decourt Quaresma.")

    figura2nome.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    figura2 = documento.add_picture(imagemDQ, width=Cm(15), height=Cm(15))

    figura2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    documento.add_page_break()

    #TABELA DECOURT QUARESMA

    dfDQ = resultadosDQ(listaSolos, estaca, diametro, listaNspt)

    tabela2nome = documento.add_paragraph("Tabela 2: Resultados pelo método de Decourt Quaresma.")

    tabela2nome.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    linhas = dfDQ.shape[0] + 1

    colunas = dfDQ.shape[1]

    tabelaDQ = documento.add_table(rows=linhas, cols=colunas, style="Light Grid Accent 1")

    for numerocoluna, indicecoluna in enumerate(dfDQ):

        for numerolinha in range(linhas):

            celula = tabelaDQ.cell(numerolinha, numerocoluna)

            if numerolinha == 0:

                celula.text = dfDQ.columns[numerocoluna]

                for paragraph in celula.paragraphs:

                    for run in paragraph.runs:

                        run.font.size = Pt(7)
            else:

                celula.text = str(dfDQ[indicecoluna][numerolinha - 1])

                for paragraph in celula.paragraphs:

                    for run in paragraph.runs:

                        run.font.size = Pt(7)

    tabelaDQ.autofit = True

    tabelaDQ.alignment = WD_TABLE_ALIGNMENT.CENTER

    documento.add_page_break()

    """"""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""

    #GRÁFICO COMPARATIVO ENTRE AOKI VELLOSO E DECOURT QUARESMA

    titulocomparativoAokieDecourt = documento.add_heading('Comparativo entre Aoki Velloso e Decourt Quaresma', level=1)

    space = documento.add_paragraph('')

    imagemAokieDecourt = plotCompararAokieDecourt(listaSolos, estaca, diametro, listaNspt, cargaAdmissivel, niveldAgua)

    figura3nome = documento.add_paragraph("Figura 3: Comparativo entre Aoki Velloso e Decourt Quaresma.")

    figura3nome.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    figura3 = documento.add_picture(imagemAokieDecourt, width=Cm(15), height=Cm(15))

    figura3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    documento.add_page_break()

    #TABELA COMPARATIVO ENTRE AOKI VELLOSO E DECOURT QUARESMA

    dfAokieDecourt = comparativoAokieDecourt(listaSolos, estaca, diametro, listaNspt)

    tabela3nome = documento.add_paragraph("Tabela 3: Comparativo entre Aoki Velloso e Decourt Quaresma.")

    tabela3nome.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    linhas = dfAokieDecourt.shape[0] + 1

    colunas = dfAokieDecourt.shape[1]

    tabelaAokieDecourt = documento.add_table(rows=linhas, cols=colunas, style="Light Grid Accent 1")

    for numerocoluna, indicecoluna in enumerate(dfAokieDecourt):

        for numerolinha in range(linhas):

            celula = tabelaAokieDecourt.cell(numerolinha, numerocoluna)

            if numerolinha == 0:

                celula.text = dfAokieDecourt.columns[numerocoluna]

                for paragraph in celula.paragraphs:

                    for run in paragraph.runs:

                        run.font.size = Pt(7)
            else:
    
                celula.text = str(dfAokieDecourt[indicecoluna][numerolinha - 1])
                
                for paragraph in celula.paragraphs:

                    for run in paragraph.runs:

                        run.font.size = Pt(7)

    tabelaAokieDecourt.autofit = True

    tabelaAokieDecourt.alignment = WD_TABLE_ALIGNMENT.CENTER

    documento.save(fileName + '.docx')

    return 
